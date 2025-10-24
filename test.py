"""
Unified NLP Processing Pipeline
Combines text preprocessing, named entity recognition, and healer scroll analysis
"""
import re
import nltk
import spacy
import pandas as pd
import logging
import json
import os
from pathlib import Path
from typing import List, Dict, Any, Union
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize, word_tokenize

# Optional file processing dependencies
try:
    import PyPDF2
    PDF_AVAILABLE = True
except:
    PDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except:
    DOCX_AVAILABLE = False

try:
    import python_docx2txt
    DOCX2TXT_AVAILABLE = True
except:
    DOCX2TXT_AVAILABLE = False

# Download necessary NLTK resources (run once)
try:
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
    nltk.download('wordnet', quiet=True)
except:
    pass

# Load spaCy model for lemmatization and NER
try:
    nlp = spacy.load('en_core_web_sm')
    SPACY_AVAILABLE = True
except:
    SPACY_AVAILABLE = False
    print("Warning: spaCy model 'en_core_web_sm' not found. NER and lemmatization will be disabled.")

# Optional dependencies for advanced features
try:
    from nltk.sentiment.vader import SentimentIntensityAnalyzer
    nltk.data.find('sentiment/vader_lexicon.zip')
    VADER_AVAILABLE = True
except:
    VADER_AVAILABLE = False

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    SKLEARN_AVAILABLE = True
except:
    SKLEARN_AVAILABLE = False

try:
    from transformers import pipeline
    TRANSFORMERS_AVAILABLE = True
except:
    TRANSFORMERS_AVAILABLE = False

logger = logging.getLogger(__name__)

# ============================================================================
# SECTION 0: FILE INPUT HANDLING
# ============================================================================

def read_pdf_file(file_path: str) -> str:
    """Extract text from PDF file using available PDF libraries."""
    text = ""
    
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip()
        except Exception as e:
            print(f"pdfplumber failed: {e}")
    
    if PDF_AVAILABLE:
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            print(f"PyPDF2 failed: {e}")
    
    raise Exception("No PDF processing library available. Install pdfplumber or PyPDF2.")

def read_json_file(file_path: str) -> str:
    """Extract text from JSON file. Handles various JSON structures."""
    with open(file_path, 'r', encoding='utf-8') as file:
        data = json.load(file)
    
    # Handle different JSON structures
    if isinstance(data, str):
        return data
    elif isinstance(data, dict):
        # Try common text fields
        text_fields = ['text', 'content', 'body', 'message', 'description', 'data']
        for field in text_fields:
            if field in data and isinstance(data[field], str):
                return data[field]
        
        # If no common field, concatenate all string values
        text_parts = []
        for key, value in data.items():
            if isinstance(value, str):
                text_parts.append(f"{key}: {value}")
            elif isinstance(value, (list, dict)):
                text_parts.append(f"{key}: {str(value)}")
        return "\n".join(text_parts)
    
    elif isinstance(data, list):
        # Handle list of strings or objects
        text_parts = []
        for item in data:
            if isinstance(item, str):
                text_parts.append(item)
            elif isinstance(item, dict):
                # Extract text from dict items
                for key, value in item.items():
                    if isinstance(value, str):
                        text_parts.append(value)
        return "\n".join(text_parts)
    
    # Fallback: convert to string
    return str(data)

def read_txt_file(file_path: str) -> str:
    """Read text from TXT file with encoding detection."""
    encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                return file.read()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            print(f"Error reading file with {encoding}: {e}")
            continue
    
    raise Exception(f"Could not read file {file_path} with any encoding")

def read_docx_file(file_path: str) -> str:
    """Extract text from Word (.docx) file using available libraries."""
    text = ""
    
    if DOCX_AVAILABLE:
        try:
            doc = Document(file_path)
            paragraphs = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            text = "\n".join(paragraphs)
            return text.strip()
            
        except Exception as e:
            print(f"python-docx failed: {e}")
    
    if DOCX2TXT_AVAILABLE:
        try:
            import docx2txt
            text = docx2txt.process(file_path)
            return text.strip()
        except Exception as e:
            print(f"docx2txt failed: {e}")
    
    raise Exception("No Word processing library available. Install python-docx or docx2txt.")

def read_input_source(input_source: Union[str, Path]) -> str:
    """
    Read text from various input sources: string, file path, or different file formats.
    
    Args:
        input_source: Can be a string (direct text), file path, or Path object
    
    Returns:
        Extracted text as string
    """
    # If it's already a string and doesn't look like a file path, return as is
    if isinstance(input_source, str) and not os.path.exists(input_source):
        # Check if it looks like a file path but doesn't exist
        if any(input_source.endswith(ext) for ext in ['.pdf', '.json', '.txt']) or '\\' in input_source or '/' in input_source:
            raise FileNotFoundError(f"File not found: {input_source}")
        return input_source
    
    # Convert to Path object for easier handling
    file_path = Path(input_source)
    
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Determine file type and process accordingly
    file_extension = file_path.suffix.lower()
    
    if file_extension == '.pdf':
        return read_pdf_file(str(file_path))
    elif file_extension == '.json':
        return read_json_file(str(file_path))
    elif file_extension in ['.txt', '.text']:
        return read_txt_file(str(file_path))
    elif file_extension in ['.docx', '.doc']:
        if file_extension == '.doc':
            raise ValueError("Legacy .doc files are not supported. Please convert to .docx format.")
        return read_docx_file(str(file_path))
    else:
        # Try to read as text file for other extensions
        try:
            return read_txt_file(str(file_path))
        except:
            raise ValueError(f"Unsupported file format: {file_extension}")

# ============================================================================
# SECTION 1: BASIC TEXT PREPROCESSING (from trial.py)
# ============================================================================

def preprocess_text(text):
    """
    Clean and normalize text using academic NLP preprocessing pipeline.
    Returns list of cleaned sentences.
    """
    if not SPACY_AVAILABLE:
        return [text.lower()]
    
    # 1) Lowercase
    text = text.lower()

    # 2) Remove symbols, numbers, punctuation
    text = re.sub(r'[^a-z\s]', ' ', text)

    # 3) Remove extra spaces
    text = re.sub(r'\s+', ' ', text).strip()

    # 4) Sentence segmentation
    sentences = sent_tokenize(text)

    # 5) Stopword removal and lemmatization per sentence
    stop_words = set(stopwords.words('english')) - {'not', 'no', 'always', 'only'}

    processed_sentences = []

    for sentence in sentences:
        words = word_tokenize(sentence)
        # Remove stopwords
        filtered_words = [word for word in words if word not in stop_words]

        # Lemmatize
        doc = nlp(" ".join(filtered_words))
        lemmatized = [token.lemma_ for token in doc]

        # Reconstruct cleaned sentence
        cleaned_sentence = ' '.join(lemmatized)
        processed_sentences.append(cleaned_sentence)

    return processed_sentences

# ============================================================================
# SECTION 2: NAMED ENTITY RECOGNITION (from NER.py)
# ============================================================================

def extract_named_entities(text):
    """
    Extract named entities from text using spaCy NLP pipeline.
    Returns list of entities with labels and positions.
    """
    if not SPACY_AVAILABLE:
        return []
    
    # Process the input text with spaCy NLP pipeline
    doc = nlp(text)

    # Extract entities and their labels
    entities = []
    for ent in doc.ents:
        entities.append({
            "text": ent.text,
            "start_char": ent.start_char,
            "end_char": ent.end_char,
            "label": ent.label_,
            "label_desc": spacy.explain(ent.label_)
        })
    return entities

# ============================================================================
# SECTION 3: HEALER SCROLL ANALYSIS (from nlp_pipeline.py)
# ============================================================================

# Healer-specific parsing constants
POSITIVE_KEYWORDS = [
    "worked", "improved", "healed", "helped", "recovered", "good", "successful", "success", "well"
]
NEGATIVE_KEYWORDS = [
    "poor", "failed", "didn't", "did not", "no help", "no improvement", "worse", "ineffective", "not help", "bad"
]

def classify_sentiment(text: str) -> str:
    """Classify sentiment of text using keyword matching."""
    t = text.lower()
    for kw in POSITIVE_KEYWORDS:
        if kw in t:
            return "positive"
    for kw in NEGATIVE_KEYWORDS:
        if kw in t:
            return "negative"
    return "neutral"

def extract_healer(text: str) -> str:
    """Extract healer name from text using regex patterns."""
    # match "Healer X" or "Healer: X" or names like "Healer Anna"
    m = re.search(r"Healer\s+([A-Z][a-zA-Z'-]+)", text)
    if m:
        return m.group(1)
    # try generic name patterns at sentence start
    m2 = re.search(r"^([A-Z][a-zA-Z'-]+)\s+used", text)
    if m2:
        return m2.group(1)
    return "Unknown"

def extract_cure_and_symptom(text: str) -> tuple:
    """Extract cure, symptom, and outcome from text using regex patterns."""
    # Try a few patterns: "used X for Y", "tried X for Y", "used X against Y"
    patterns = [
        r"used\s+([a-zA-Z0-9\s'-]+?)\s+for\s+([a-zA-Z0-9\s'-]+)[,\.-]?(.*)$",
        r"tried\s+([a-zA-Z0-9\s'-]+?)\s+for\s+([a-zA-Z0-9\s'-]+)[,\.-]?(.*)$",
        r"used\s+([a-zA-Z0-9\s'-]+?)\s+against\s+([a-zA-Z0-9\s'-]+)[,\.-]?(.*)$",
        r"applied\s+([a-zA-Z0-9\s'-]+?)\s+for\s+([a-zA-Z0-9\s'-]+)[,\.-]?(.*)$",
    ]
    for p in patterns:
        m = re.search(p, text, flags=re.IGNORECASE)
        if m:
            cure = m.group(1).strip()
            symptom = m.group(2).strip()
            outcome = m.group(3).strip()
            return cure, symptom, outcome

    # fallback: try to find the word after 'used' and optionally a following 'for'
    m2 = re.search(r"used\s+([a-zA-Z0-9\s'-]+)", text, flags=re.IGNORECASE)
    if m2:
        cure = m2.group(1).strip()
        m3 = re.search(r"for\s+([a-zA-Z0-9\s'-]+)", text, flags=re.IGNORECASE)
        if m3:
            symptom = m3.group(1).strip()
            rest = text[m3.end():].strip()
            return cure, symptom, rest
        else:
            return cure, "unknown", ""

    return "", "", ""

def parse_text(text: str) -> List[Dict]:
    """Parse healer text and extract structured information."""
    sentences = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]
    
    results = []
    for sentence in sentences:
        if len(sentence) < 10:
            continue
            
        healer = extract_healer(sentence)
        cure, symptom, outcome = extract_cure_and_symptom(sentence)
        sentiment = classify_sentiment(sentence)
        
        record = {
            'healer': healer,
            'cure': cure,
            'symptom': symptom,
            'outcome': outcome,
            'sentiment': sentiment,
            'raw': sentence
        }
        results.append(record)
    
    return results

def clean_text(text: str) -> str:
    """Clean text by normalizing whitespace."""
    return re.sub(r"\s+", " ", text.replace('\r', ' ')).strip()

def extract_keywords_tfidf(texts: List[str], top_n: int = 10) -> List[str]:
    """Extract keywords using TF-IDF or fallback to frequency counting."""
    if not SKLEARN_AVAILABLE:
        # fallback to simple frequency
        words = ' '.join(texts).lower().split()
        freq = {}
        for w in words:
            if len(w) < 3:
                continue
            freq[w] = freq.get(w, 0) + 1
        return [k for k, _ in sorted(freq.items(), key=lambda x: -x[1])][:top_n]

    vectorizer = TfidfVectorizer(stop_words='english', max_features=2000)
    X = vectorizer.fit_transform(texts)
    scores = X.sum(axis=0).A1
    indices = scores.argsort()[-top_n:][::-1]
    return [vectorizer.get_feature_names_out()[i] for i in indices]

def summarize_with_transformer(text: str) -> str:
    """Generate summary using transformer model with fallback."""
    if not TRANSFORMERS_AVAILABLE:
        return ""  # caller will handle fallback
    try:
        summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")
        res = summarizer(text, max_length=120, min_length=20, do_sample=False)
        return res[0]['summary_text']
    except Exception as e:
        logger.warning("Transformer summarization failed: %s", e)
        return ""

def analyze_sentiments_vader(outcomes: List[str]) -> List[float]:
    """Analyze sentiment scores using VADER or fallback method."""
    if not VADER_AVAILABLE:
        # fallback: map 'positive'/'negative' strings if present
        scores = []
        for o in outcomes:
            s = o.lower()
            if any(k in s for k in ['work', 'improv', 'heal', 'help']):
                scores.append(0.6)
            elif any(k in s for k in ['poor', 'fail', "didn't", 'did not', 'no help', 'worse']):
                scores.append(-0.6)
            else:
                scores.append(0.0)
        return scores

    sid = SentimentIntensityAnalyzer()
    return [sid.polarity_scores(o)['compound'] for o in outcomes]

def process_scrolls(text: str) -> Dict[str, Any]:
    """Process raw healer scrolls and return structured insights."""
    text = clean_text(text)
    records = parse_text(text)
    df = pd.DataFrame(records)

    # compute counts
    cures_pos = {}
    cures_neg = {}
    if not df.empty:
        for _, r in df.iterrows():
            cure = (r.get('cure') or '').strip()
            s = r.get('sentiment')
            if not cure:
                continue
            if s == 'positive':
                cures_pos[cure] = cures_pos.get(cure, 0) + 1
            elif s == 'negative':
                cures_neg[cure] = cures_neg.get(cure, 0) + 1

    # keywords from cures/symptoms/raw
    texts_for_k = []
    if not df.empty:
        texts_for_k = (df['cure'].fillna('') + ' ' + df['symptom'].fillna('') + ' ' + df['raw'].fillna('')).tolist()
    else:
        texts_for_k = [text]

    keywords = extract_keywords_tfidf(texts_for_k, top_n=12)

    # sentiment scores for outcomes
    outcomes = df['outcome'].fillna('').tolist() if not df.empty else [text]
    sentiment_scores = analyze_sentiments_vader(outcomes)

    # try transformer summarization
    summary = ''
    if TRANSFORMERS_AVAILABLE:
        summary = summarize_with_transformer(text)

    if not summary:
        # rule-based summary: list top positive and negative cures
        pos_sorted = [k for k, _ in sorted(cures_pos.items(), key=lambda x: -x[1])]
        neg_sorted = [k for k, _ in sorted(cures_neg.items(), key=lambda x: -x[1])]
        parts = []
        if pos_sorted:
            parts.append(f"Frequent effective cures: {', '.join(pos_sorted[:5])}.")
        if neg_sorted:
            parts.append(f"Frequent ineffective cures: {', '.join(neg_sorted[:5])}.")
        if not parts:
            parts = ["No clear wisdom extracted — add more notes or enable transformer summarization."]
        summary = ' '.join(parts)

    result = {
        'records': records,
        'cures_pos_counts': cures_pos,
        'cures_neg_counts': cures_neg,
        'keywords': keywords,
        'summary': summary,
        'sentiment_scores': sentiment_scores,
    }

    return result

# ============================================================================
# SECTION 4: UNIFIED PROCESSING FUNCTION
# ============================================================================

def process_text_complete(input_source: Union[str, Path], include_preprocessing=True, include_ner=True, include_healer_analysis=True):
    """
    Complete text processing pipeline combining all three approaches.
    
    Args:
        input_source: Input text string or file path (supports .txt, .pdf, .json)
        include_preprocessing: Whether to run basic text preprocessing
        include_ner: Whether to extract named entities
        include_healer_analysis: Whether to run healer scroll analysis
    
    Returns:
        Dictionary containing results from all enabled processing steps
    """
    # Extract text from input source (file or string)
    text = read_input_source(input_source)
    
    # Determine input type for metadata
    input_type = "string"
    if isinstance(input_source, (str, Path)) and os.path.exists(str(input_source)):
        input_type = f"file ({Path(input_source).suffix})"
    
    results = {
        'input_source': str(input_source),
        'input_type': input_type,
        'original_text': text,
        'text_length': len(text),
        'preprocessing_available': SPACY_AVAILABLE,
        'ner_available': SPACY_AVAILABLE,
        'file_processing_available': {
            'pdf_pdfplumber': PDFPLUMBER_AVAILABLE,
            'pdf_pypdf2': PDF_AVAILABLE,
            'docx_python_docx': DOCX_AVAILABLE,
            'docx_docx2txt': DOCX2TXT_AVAILABLE,
            'json': True,
            'txt': True
        },
        'advanced_features_available': {
            'vader_sentiment': VADER_AVAILABLE,
            'sklearn_tfidf': SKLEARN_AVAILABLE,
            'transformers': TRANSFORMERS_AVAILABLE
        }
    }
    
    if include_preprocessing:
        results['preprocessed_sentences'] = preprocess_text(text)
    
    if include_ner:
        results['named_entities'] = extract_named_entities(text)
    
    if include_healer_analysis:
        results['healer_analysis'] = process_scrolls(text)
    
    return results

# ============================================================================
# SECTION 5: EXAMPLE USAGE AND TESTING
# ============================================================================

def process_file_or_text(input_source: Union[str, Path], output_format="summary"):
    """
    Convenience function to process any input and return results in different formats.
    
    Args:
        input_source: Text string or file path (supports .txt, .pdf, .json)
        output_format: "summary", "detailed", or "raw"
    
    Returns:
        Processed results in the specified format
    """
    try:
        results = process_text_complete(input_source)
        
        if output_format == "raw":
            return results
        
        elif output_format == "detailed":
            return {
                'input_info': {
                    'source': results.get('input_source'),
                    'type': results.get('input_type'),
                    'length': results.get('text_length')
                },
                'preprocessing': {
                    'sentences': results.get('preprocessed_sentences', []),
                    'sentence_count': len(results.get('preprocessed_sentences', []))
                },
                'entities': {
                    'list': results.get('named_entities', []),
                    'count': len(results.get('named_entities', []))
                },
                'healer_analysis': results.get('healer_analysis', {}),
                'capabilities': results.get('file_processing_available', {})
            }
        
        else:  # summary format
            healer_data = results.get('healer_analysis', {})
            return {
                'input_type': results.get('input_type'),
                'text_length': results.get('text_length'),
                'entities_found': len(results.get('named_entities', [])),
                'healer_records': len(healer_data.get('records', [])),
                'positive_cures': healer_data.get('cures_pos_counts', {}),
                'negative_cures': healer_data.get('cures_neg_counts', {}),
                'summary': healer_data.get('summary', ''),
                'top_keywords': healer_data.get('keywords', [])[:5]
            }
    
    except Exception as e:
        return {'error': str(e)}

def save_results_to_json(all_results: List[Dict], output_filename: str = "output.json"):
    """
    Save all processing results to a JSON file.
    
    Args:
        all_results: List of dictionaries containing processing results
        output_filename: Name of the output JSON file
    """
    try:
        # Add metadata to the output
        output_data = {
            "metadata": {
                "processing_date": pd.Timestamp.now().isoformat(),
                "total_inputs_processed": len(all_results),
                "nlp_pipeline_version": "1.0.0",
                "system_capabilities": {
                    "spacy_available": SPACY_AVAILABLE,
                    "vader_available": VADER_AVAILABLE,
                    "sklearn_available": SKLEARN_AVAILABLE,
                    "transformers_available": TRANSFORMERS_AVAILABLE,
                    "pdf_processing": {
                        "pdfplumber": PDFPLUMBER_AVAILABLE,
                        "pypdf2": PDF_AVAILABLE
                    },
                    "word_processing": {
                        "python_docx": DOCX_AVAILABLE,
                        "docx2txt": DOCX2TXT_AVAILABLE
                    }
                }
            },
            "results": all_results
        }
        
        # Save to JSON file with proper formatting
        with open(output_filename, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, indent=2, ensure_ascii=False, default=str)
        
        print(f"\n✅ Results saved to: {output_filename}")
        print(f"   File size: {os.path.getsize(output_filename)} bytes")
        return True
        
    except Exception as e:
        print(f"❌ Error saving results to JSON: {str(e)}")
        return False

if __name__ == "__main__":
    print("=" * 80)
    print("UNIFIED NLP PROCESSING PIPELINE - MULTI-FORMAT INPUT")
    print("=" * 80)
    
    
    
    
    # List to store all processing results for JSON output
    all_processing_results = []
    
    
    
    input_names = ["Direct String", "JSON File", "TXT File", "Word File", "PDF File"]
    
    for i, (input_source, input_name) in enumerate(zip(test_inputs, input_names), 1):
        print(f"\n{'='*80}")
        print(f"TEST {i}: {input_name.upper()} INPUT")
        print(f"{'='*80}")
        
        try:
            # Read and display source
            text = read_input_source(input_source)
            print(f"INPUT SOURCE: {input_source}")
            print(f"INPUT TYPE: {input_name}")
            print(f"TEXT LENGTH: {len(text)} characters")
            print(f"TEXT PREVIEW: {text[:200]}{'...' if len(text) > 200 else ''}")
            
            print(f"\n{'-'*60}")
            
            # Process 1: Basic text preprocessing
            print("1. BASIC TEXT PREPROCESSING:")
            preprocessed = preprocess_text(text)
            print(f"   Preprocessed into {len(preprocessed)} sentences")
            for j, sentence in enumerate(preprocessed[:3], 1):  # Show first 3
                print(f"   {j}. {sentence[:80]}{'...' if len(sentence) > 80 else ''}")
            if len(preprocessed) > 3:
                print(f"   ... and {len(preprocessed) - 3} more sentences")
            
            # Process 2: Named Entity Recognition
            print("\n2. NAMED ENTITY RECOGNITION:")
            entities = extract_named_entities(text)
            if entities:
                print(f"   Found {len(entities)} named entities:")
                for e in entities[:5]:  # Show first 5
                    print(f"   • {e['text']} ({e['label']} - {e['label_desc']})")
                if len(entities) > 5:
                    print(f"   ... and {len(entities) - 5} more entities")
            else:
                print("   No named entities found or spaCy not available.")
            
            # Process 3: Healer scroll analysis
            print("\n3. HEALER SCROLL ANALYSIS:")
            healer_results = process_scrolls(text)
            print(f"   Extracted {len(healer_results['records'])} healer records:")
            for j, record in enumerate(healer_results['records'][:3], 1):  # Show first 3
                print(f"   Record {j}: {record['healer']} used {record['cure']} for {record['symptom']} ({record['sentiment']})")
            if len(healer_results['records']) > 3:
                print(f"   ... and {len(healer_results['records']) - 3} more records")
            
            print(f"\n   Summary: {healer_results['summary']}")
            
            # Process 4: Complete unified processing
            print("\n4. COMPLETE UNIFIED PROCESSING:")
            complete_results = process_text_complete(input_source)
            print(f"   Input Type: {complete_results.get('input_type', 'Unknown')}")
            print(f"   Text Length: {complete_results.get('text_length', 0)} characters")
            print(f"   Preprocessed sentences: {len(complete_results.get('preprocessed_sentences', []))}")
            print(f"   Named entities: {len(complete_results.get('named_entities', []))}")
            print(f"   Healer records: {len(complete_results.get('healer_analysis', {}).get('records', []))}")
            
            # Store results for JSON output
            result_entry = {
                "test_number": i,
                "input_name": input_name,
                "input_source": str(input_source),
                "processing_results": complete_results,
                "individual_processing": {
                    "preprocessing": {
                        "sentences": preprocessed,
                        "sentence_count": len(preprocessed)
                    },
                    "named_entities": {
                        "entities": entities,
                        "entity_count": len(entities)
                    },
                    "healer_analysis": healer_results
                }
            }
            all_processing_results.append(result_entry)
            
        except Exception as e:
            print(f"ERROR processing {input_name}: {str(e)}")
            # Store error result
            error_entry = {
                "test_number": i,
                "input_name": input_name,
                "input_source": str(input_source),
                "error": str(e),
                "processing_results": None
            }
            all_processing_results.append(error_entry)
            continue
    
    # Show system capabilities
    print(f"\n{'='*80}")
    print("SYSTEM CAPABILITIES")
    print(f"{'='*80}")
    print("Text Processing:")
    print(f"  ✓ spaCy (preprocessing/NER): {'Available' if SPACY_AVAILABLE else 'Not Available'}")
    print(f"  ✓ VADER sentiment: {'Available' if VADER_AVAILABLE else 'Not Available'}")
    print(f"  ✓ scikit-learn TF-IDF: {'Available' if SKLEARN_AVAILABLE else 'Not Available'}")
    print(f"  ✓ Transformers: {'Available' if TRANSFORMERS_AVAILABLE else 'Not Available'}")
    
    print("\nFile Processing:")
    print(f"  ✓ JSON files: Available")
    print(f"  ✓ TXT files: Available")
    print(f"  ✓ Word files (python-docx): {'Available' if DOCX_AVAILABLE else 'Not Available'}")
    print(f"  ✓ Word files (docx2txt): {'Available' if DOCX2TXT_AVAILABLE else 'Not Available'}")
    print(f"  ✓ PDF files (pdfplumber): {'Available' if PDFPLUMBER_AVAILABLE else 'Not Available'}")
    print(f"  ✓ PDF files (PyPDF2): {'Available' if PDF_AVAILABLE else 'Not Available'}")
    
    print(f"\nSupported Input Formats:")
    print(f"  • Direct text strings")
    print(f"  • .txt and .text files")
    print(f"  • .json files (flexible structure)")
    print(f"  • .docx files (Word documents - if libraries installed)")
    print(f"  • .pdf files (if PDF libraries installed)")
    
    print(f"\n{'='*80}")
    print("PROCESSING COMPLETE - Tested multiple input formats")
    print(f"{'='*80}")
    
    # Save all results to JSON file
    print(f"\n{'='*80}")
    print("SAVING RESULTS TO JSON")
    print(f"{'='*80}")
    
    if all_processing_results:
        success = save_results_to_json(all_processing_results, "output.json")
        if success:
            print(f"📄 All processing results have been saved to 'output.json'")
            print(f"   Total inputs processed: {len(all_processing_results)}")
            print(f"   You can open this file to view detailed results")
        else:
            print("❌ Failed to save results to JSON file")
    else:
        print("⚠️  No results to save")
    
    # Clean up sample files
    try:
        os.remove('sample_healer_data.json')
        os.remove('sample_healing_records.txt')
        if os.path.exists('sample_medieval_chronicles.docx'):
            os.remove('sample_medieval_chronicles.docx')
        print("\nSample files cleaned up.")
    except:
        pass
    
    print(f"\n{'='*80}")
    print("🎉 ALL PROCESSING AND SAVING COMPLETE!")
    print(f"{'='*80}")